import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
import whisper
from tkinterdnd2 import TkinterDnD, DND_FILES
from moviepy import VideoFileClip
from PIL import Image, ImageTk
import threading
import os
import sys
import time
from datetime import datetime

# Global variables
last_transcription_path = None
transcribed_text = ""
current_file_size = 0
processing_start_time = None


class ModernProgressBar:
    def __init__(self, parent):
        self.parent = parent
        self.frame = tk.Frame(parent, bg="#f8f9fa")

        # Progress bar container
        self.progress_container = tk.Frame(self.frame, bg="#f8f9fa", relief="flat", bd=0)
        self.progress_container.pack(fill="x", padx=20, pady=10)

        # Progress bar
        style = ttk.Style()
        style.configure("Modern.Horizontal.TProgressbar",
                        troughcolor="#e9ecef",
                        background="#007bff",
                        bordercolor="#e9ecef",
                        lightcolor="#007bff",
                        darkcolor="#007bff")

        self.progress_bar = ttk.Progressbar(
            self.progress_container,
            orient="horizontal",
            length=400,
            mode="determinate",
            style="Modern.Horizontal.TProgressbar"
        )
        self.progress_bar.pack(fill="x", pady=(0, 5))

        # Progress text
        self.progress_text = tk.Label(
            self.progress_container,
            text="0%",
            font=("Segoe UI", 10, "bold"),
            bg="#f8f9fa",
            fg="#007bff"
        )
        self.progress_text.pack()

        # Status text
        self.status_text = tk.Label(
            self.progress_container,
            text="Ready to process",
            font=("Segoe UI", 9),
            bg="#f8f9fa",
            fg="#6c757d"
        )
        self.status_text.pack(pady=(2, 0))

        # Time elapsed
        self.time_text = tk.Label(
            self.progress_container,
            text="",
            font=("Segoe UI", 8),
            bg="#f8f9fa",
            fg="#adb5bd"
        )
        self.time_text.pack()

        self.start_time = None

    def update_progress(self, percent, status="Processing..."):
        self.progress_bar["value"] = percent
        self.progress_text.config(text=f"{percent}%")
        self.status_text.config(text=status)

        if self.start_time and percent > 0:
            elapsed = time.time() - self.start_time
            self.time_text.config(text=f"Elapsed: {elapsed:.1f}s")

        self.parent.update_idletasks()

    def start_progress(self):
        self.start_time = time.time()
        self.progress_bar["value"] = 0
        self.progress_text.config(text="0%")
        self.status_text.config(text="Starting...")
        self.time_text.config(text="")
        self.parent.update_idletasks()

    def reset_progress(self):
        self.progress_bar["value"] = 0
        self.progress_text.config(text="0%")
        self.status_text.config(text="Ready to process")
        self.time_text.config(text="")
        self.start_time = None
        self.parent.update_idletasks()


def extract_audio(video_path, output_audio_path, progress_callback):
    progress_callback(10, "Loading video file...")
    clip = VideoFileClip(video_path)

    progress_callback(20, "Extracting audio...")
    clip.audio.write_audiofile(output_audio_path, logger=None)

    progress_callback(30, "Audio extraction complete")
    clip.close()


def export_srt(segments, path):
    def format_timestamp(seconds):
        h = int(seconds // 3600)
        m = int((seconds % 3600) // 60)
        s = int(seconds % 60)
        ms = int((seconds - int(seconds)) * 1000)
        return f"{h:02}:{m:02}:{s:02},{ms:03}"

    with open(path, "w", encoding="utf-8") as f:
        for i, seg in enumerate(segments, 1):
            start = format_timestamp(seg["start"])
            end = format_timestamp(seg["end"])
            text = seg["text"].strip()
            f.write(f"{i}\n{start} --> {end}\n{text}\n\n")


def export_txt(segments, path):
    with open(path, "w", encoding="utf-8") as f:
        for seg in segments:
            text = seg["text"].strip()
            f.write(f"{text}\n")


def export_pdf(segments, path):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        doc = SimpleDocTemplate(path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center
        )
        story.append(Paragraph("Transcription Results", title_style))
        story.append(Spacer(1, 12))
        
        # Content
        for i, seg in enumerate(segments, 1):
            text = seg["text"].strip()
            if text:
                story.append(Paragraph(f"{text}", styles['Normal']))
                story.append(Spacer(1, 6))
        
        doc.build(story)
        return True
    except ImportError:
        return False


def export_word(segments, path):
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Transcription Results', 0)
        title.alignment = 1  # Center
        
        # Content
        for seg in segments:
            text = seg["text"].strip()
            if text:
                doc.add_paragraph(text)
        
        doc.save(path)
        return True
    except ImportError:
        return False


def export_file(segments, path, format_type):
    """Export transcription to the specified format"""
    if format_type == "SRT":
        export_srt(segments, path)
    elif format_type == "TXT":
        export_txt(segments, path)
    elif format_type == "PDF":
        if not export_pdf(segments, path):
            raise Exception("PDF export requires reportlab library. Install with: pip install reportlab")
    elif format_type == "Word":
        if not export_word(segments, path):
            raise Exception("Word export requires python-docx library. Install with: pip install python-docx")
    else:
        raise Exception(f"Unsupported format: {format_type}")


def transcribe_audio(audio_path, progress_callback):
    progress_callback(40, "Loading AI model...")
    model = whisper.load_model("base")

    progress_callback(50, "Transcribing audio...")
    result = model.transcribe(audio_path, language="english", word_timestamps=True)

    progress_callback(80, "Transcription complete")
    return result


def update_progress_callback(percent, status):
    if hasattr(root, 'progress_bar'):
        root.progress_bar.update_progress(percent, status)


def process_path(file_path):
    global last_transcription_path, transcribed_text

    if not file_path:
        return

    # Update file info
    file_name = os.path.basename(file_path)
    file_size = os.path.getsize(file_path) / (1024 * 1024)  # MB

    root.file_label.config(text=f"📁 {file_name}")
    root.file_size_label.config(text=f"Size: {file_size:.1f} MB")

    # Start progress
    root.progress_bar.start_progress()

    is_video = file_path.lower().endswith((".mp4", ".mkv", ".avi"))
    audio_path = "temp_audio.wav" if is_video else file_path

    try:
        if is_video:
            extract_audio(file_path, audio_path, update_progress_callback)
        else:
            update_progress_callback(30, "Audio file ready")

        result = transcribe_audio(audio_path, update_progress_callback)

        update_progress_callback(90, "Saving results...")

        # Get selected format
        selected_format = root.format_var.get()
        
        # Extract format type from dropdown value
        if "SRT" in selected_format:
            format_type = "SRT"
            defaultextension = ".srt"
            filetypes = [("SubRip Subtitle", "*.srt")]
        elif "TXT" in selected_format:
            format_type = "TXT"
            defaultextension = ".txt"
            filetypes = [("Text files", "*.txt")]
        elif "PDF" in selected_format:
            format_type = "PDF"
            defaultextension = ".pdf"
            filetypes = [("PDF files", "*.pdf")]
        elif "Word" in selected_format:
            format_type = "Word"
            defaultextension = ".docx"
            filetypes = [("Word documents", "*.docx")]
        else:
            format_type = "SRT"
            defaultextension = ".srt"
            filetypes = [("SubRip Subtitle", "*.srt")]
        
        save_path = filedialog.asksaveasfilename(
            defaultextension=defaultextension,
            filetypes=filetypes
        )

        if save_path:
            export_file(result["segments"], save_path, format_type)
            last_transcription_path = save_path
            update_progress_callback(100, "Transcription saved successfully!")

            # Show success message
            messagebox.showinfo("Success",
                                f"Transcription completed successfully!\nSaved to: {os.path.basename(save_path)}")

            # Reset after delay
            root.after(2000, root.progress_bar.reset_progress)

        if is_video and os.path.exists("temp_audio.wav"):
            os.remove("temp_audio.wav")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred:\n{str(e)}")
        root.progress_bar.reset_progress()


def select_file():
    file_path = filedialog.askopenfilename(
        filetypes=[("Audio/Video Files", "*.mp3 *.mp4 *.wav *.m4a *.mkv *.avi")]
    )
    if file_path:
        threading.Thread(target=process_path, args=(file_path,), daemon=True).start()


def drop(event):
    file_path = event.data.strip("{}")
    if os.path.isfile(file_path):
        threading.Thread(target=process_path, args=(file_path,), daemon=True).start()


def create_modern_button(parent, text, command, bg_color="#007bff", hover_color="#0056b3"):
    button_frame = tk.Frame(parent, bg="#f8f9fa")

    def on_enter(e):
        button.config(bg=hover_color)

    def on_leave(e):
        button.config(bg=bg_color)

    button = tk.Button(
        button_frame,
        text=text,
        command=command,
        font=("Segoe UI", 11, "bold"),
        bg=bg_color,
        fg="white",
        relief="flat",
        bd=0,
        padx=30,
        pady=12,
        cursor="hand2"
    )
    button.pack()
    button.bind("<Enter>", on_enter)
    button.bind("<Leave>", on_leave)

    return button_frame


# --- GUI setup ---
root = TkinterDnD.Tk()


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


# Set window properties
root.title("Aurora Speech Extractor")
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()
root.geometry(f"{screen_width}x{screen_height}")
root.configure(bg="#f8f9fa")
root.resizable(True, True)

# Set icon
icon_path = resource_path("icon.ico")
if os.path.exists(icon_path):
    root.iconbitmap(default=icon_path)

# Main container
main_container = tk.Frame(root, bg="#f8f9fa")
main_container.pack(fill="both", expand=True, padx=20, pady=20)

# Header with centered logo above text
header_frame = tk.Frame(main_container, bg="#f8f9fa")
header_frame.pack(fill="x", pady=(5, 10))

# Logo centered above text
try:
    logo_path = resource_path("logo.png")
    logo_image = Image.open(logo_path).resize((120, 120), Image.Resampling.LANCZOS)
    logo_photo = ImageTk.PhotoImage(logo_image)
    logo_label = tk.Label(header_frame, image=logo_photo, bg="#f8f9fa")
    logo_label.image = logo_photo
    logo_label.pack(expand=True, pady=(0, 15))
except Exception as e:
    print("Logo not loaded:", e)

# Text container (centered below logo)
text_container = tk.Frame(header_frame, bg="#f8f9fa")
text_container.pack(expand=True)

# App description
app_description = tk.Label(
    text_container,
    text="AURORA - AI Speech to Text Converter",
    font=("Segoe UI", 18, "bold"),
    bg="#f8f9fa",
    fg="#212529"
)
app_description.pack()

sub_description = tk.Label(
    text_container,
    text="Convert audio and video files to text using advanced AI technology",
    font=("Segoe UI", 12),
    bg="#f8f9fa",
    fg="#6c757d"
)
sub_description.pack(pady=(5, 0))

# Drag and drop area
drop_frame = tk.Frame(main_container, bg="#ffffff", relief="solid", bd=1)
drop_frame.pack(fill="x", pady=(0, 30))

drop_label = tk.Label(
    drop_frame,
    text="📁 Drop your audio/video file here",
    font=("Segoe UI", 14),
    bg="#ffffff",
    fg="#6c757d",
    pady=40
)
drop_label.pack()

# Select button inside drop area
select_button = create_modern_button(drop_frame, "📂 Select File", select_file)
select_button.pack(pady=(0, 40))

# Format selection
format_frame = tk.Frame(main_container, bg="#f8f9fa")
format_frame.pack(fill="x", pady=(0, 20))

tk.Label(
    format_frame,
    text="Output Format:",
    font=("Segoe UI", 11, "bold"),
    bg="#f8f9fa",
    fg="#212529"
).pack(anchor="w", pady=(0, 5))

# Format dropdown
root.format_var = tk.StringVar(value="SRT")
format_options = ["SRT (Subtitles)", "TXT (Plain Text)", "PDF", "Word (.docx)"]
format_dropdown = ttk.Combobox(
    format_frame,
    textvariable=root.format_var,
    values=format_options,
    state="readonly",
    font=("Segoe UI", 10),
    width=25
)
format_dropdown.pack(anchor="w")

# File info
file_info_frame = tk.Frame(main_container, bg="#f8f9fa")
file_info_frame.pack(fill="x", pady=(0, 20))

root.file_label = tk.Label(
    file_info_frame,
    text="No file selected",
    font=("Segoe UI", 11),
    bg="#f8f9fa",
    fg="#6c757d"
)
root.file_label.pack()

root.file_size_label = tk.Label(
    file_info_frame,
    text="",
    font=("Segoe UI", 9),
    bg="#f8f9fa",
    fg="#adb5bd"
)
root.file_size_label.pack()

# Progress section
progress_frame = tk.Frame(main_container, bg="#f8f9fa")
progress_frame.pack(fill="x", pady=(0, 20))

tk.Label(
    progress_frame,
    text="Processing Progress",
    font=("Segoe UI", 14, "bold"),
    bg="#f8f9fa",
    fg="#212529"
).pack(anchor="w", pady=(0, 10))

root.progress_bar = ModernProgressBar(progress_frame)
root.progress_bar.frame.pack(fill="x")

# Info section
info_frame = tk.Frame(main_container, bg="#e9ecef", relief="flat", bd=0)
info_frame.pack(fill="x", pady=(20, 0))

info_text = (
    "💡 Tips:\n"
    "• Use clear audio/video with minimal background noise\n"
    "• Supported formats: MP3, WAV, MP4, MKV, AVI, M4A\n"
    "• Processing time depends on file size and system performance\n"
    "• Results are saved as SRT subtitle files"
)

tk.Label(
    info_frame,
    text=info_text,
    font=("Segoe UI", 10),
    bg="#e9ecef",
    fg="#495057",
    justify="left",
    wraplength=800
).pack(padx=20, pady=15)

# Footer
footer_frame = tk.Frame(main_container, bg="#f8f9fa")
footer_frame.pack(fill="x", pady=(20, 0))

footer = tk.Label(
    footer_frame,
    text="© 2025 Pynexor Ltd. · All rights reserved.",
    font=("Segoe UI", 9),
    bg="#f8f9fa",
    fg="#adb5bd"
)
footer.pack(side="bottom")

# Setup drag and drop
root.drop_target_register(DND_FILES)
root.dnd_bind('<<Drop>>', drop)

# Make drop area clickable
drop_label.bind("<Button-1>", lambda e: select_file())
drop_frame.bind("<Button-1>", lambda e: select_file())

root.mainloop()
